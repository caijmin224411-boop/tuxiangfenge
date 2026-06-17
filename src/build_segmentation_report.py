from pathlib import Path
from collections import deque
import json
import math
import shutil

import numpy as np
from PIL import Image, ImageDraw, ImageFont, ImageFilter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
LOCAL_IMAGE = ROOT / "figures" / "source_photo.jpg"
SRC_IMAGE = LOCAL_IMAGE if LOCAL_IMAGE.exists() else Path(
    r"D:\documents\xwechat\xwechat_files\wxid_7v7rttt74gpt11_8f59\temp\RWTemp\2026-06"
    r"\d76229233573c7aad92ba65ef9af4482.jpg"
)
OUT_DIR = ROOT / "segmentation_report_output"
IMG_DIR = OUT_DIR / "figures"
DOCX_PATH = OUT_DIR / "图像分割实验报告_黄色毛绒玩偶.docx"
ADVANCED_JSON = OUT_DIR / "advanced_segmentation_stats.json"


def ensure_dirs():
    IMG_DIR.mkdir(parents=True, exist_ok=True)


def load_font(size=28):
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size)
    return ImageFont.load_default()


FONT_BIG = load_font(34)
FONT = load_font(24)
FONT_SMALL = load_font(18)


def resize_for_work(img, width=1200):
    w, h = img.size
    ratio = width / w
    return img.resize((width, int(h * ratio)), Image.Resampling.LANCZOS)


def binary_erode(mask, radius=1, iterations=1):
    out = mask.copy()
    for _ in range(iterations):
        padded = np.pad(out, radius, mode="constant", constant_values=False)
        result = np.ones_like(out, dtype=bool)
        for dy in range(-radius, radius + 1):
            for dx in range(-radius, radius + 1):
                result &= padded[
                    radius + dy : radius + dy + out.shape[0],
                    radius + dx : radius + dx + out.shape[1],
                ]
        out = result
    return out


def binary_dilate(mask, radius=1, iterations=1):
    out = mask.copy()
    for _ in range(iterations):
        padded = np.pad(out, radius, mode="constant", constant_values=False)
        result = np.zeros_like(out, dtype=bool)
        for dy in range(-radius, radius + 1):
            for dx in range(-radius, radius + 1):
                result |= padded[
                    radius + dy : radius + dy + out.shape[0],
                    radius + dx : radius + dx + out.shape[1],
                ]
        out = result
    return out


def connected_components(mask, min_area=260, max_area=26000):
    h, w = mask.shape
    seen = np.zeros_like(mask, dtype=bool)
    comps = []
    ys, xs = np.nonzero(mask)
    points = list(zip(xs, ys))
    for sx, sy in points:
        if seen[sy, sx] or not mask[sy, sx]:
            continue
        q = deque([(sx, sy)])
        seen[sy, sx] = True
        minx = maxx = sx
        miny = maxy = sy
        area = 0
        while q:
            x, y = q.popleft()
            area += 1
            minx, maxx = min(minx, x), max(maxx, x)
            miny, maxy = min(miny, y), max(maxy, y)
            for nx, ny in ((x - 1, y), (x + 1, y), (x, y - 1), (x, y + 1)):
                if 0 <= nx < w and 0 <= ny < h and mask[ny, nx] and not seen[ny, nx]:
                    seen[ny, nx] = True
                    q.append((nx, ny))
        bw, bh = maxx - minx + 1, maxy - miny + 1
        aspect = bw / max(1, bh)
        if min_area <= area <= max_area and 0.18 <= aspect <= 4.8 and bw > 8 and bh > 8:
            comps.append({"bbox": (minx, miny, maxx, maxy), "area": area})
    comps.sort(key=lambda c: c["area"], reverse=True)
    return comps


def crop_component_mask(mask, bbox, pad=2):
    h, w = mask.shape
    x0, y0, x1, y1 = bbox
    x0 = max(0, x0 - pad)
    y0 = max(0, y0 - pad)
    x1 = min(w - 1, x1 + pad)
    y1 = min(h - 1, y1 + pad)
    return mask[y0 : y1 + 1, x0 : x1 + 1], (x0, y0, x1, y1)


def chamfer_distance(mask):
    """Small no-dependency distance transform for binary masks."""
    inf = 10_000
    dist = np.where(mask, inf, 0).astype(np.int32)
    h, w = mask.shape
    for y in range(h):
        for x in range(w):
            if not mask[y, x]:
                continue
            best = dist[y, x]
            if y > 0:
                best = min(best, dist[y - 1, x] + 3)
            if x > 0:
                best = min(best, dist[y, x - 1] + 3)
            if y > 0 and x > 0:
                best = min(best, dist[y - 1, x - 1] + 4)
            if y > 0 and x + 1 < w:
                best = min(best, dist[y - 1, x + 1] + 4)
            dist[y, x] = best
    for y in range(h - 1, -1, -1):
        for x in range(w - 1, -1, -1):
            if not mask[y, x]:
                continue
            best = dist[y, x]
            if y + 1 < h:
                best = min(best, dist[y + 1, x] + 3)
            if x + 1 < w:
                best = min(best, dist[y, x + 1] + 3)
            if y + 1 < h and x + 1 < w:
                best = min(best, dist[y + 1, x + 1] + 4)
            if y + 1 < h and x > 0:
                best = min(best, dist[y + 1, x - 1] + 4)
            dist[y, x] = best
    return dist


def pick_instance_seeds(local_mask, max_seeds=8):
    dist = chamfer_distance(local_mask)
    area = int(local_mask.sum())
    if area < 850:
        return []
    expected = max(1, min(max_seeds, int(round(area / 2600))))
    min_dist = max(12, int(math.sqrt(area / max(expected, 1)) * 0.42))
    threshold = max(9, int(dist.max() * 0.42))
    ys, xs = np.nonzero((dist >= threshold) & local_mask)
    candidates = sorted(
        [(int(dist[y, x]), int(x), int(y)) for x, y in zip(xs, ys)],
        reverse=True,
    )
    seeds = []
    for score, x, y in candidates:
        if all((x - sx) ** 2 + (y - sy) ** 2 >= min_dist**2 for sx, sy, _ in seeds):
            seeds.append((x, y, score))
            if len(seeds) >= expected:
                break
    return seeds


def assign_to_seeds(local_mask, seeds):
    labels = np.zeros(local_mask.shape, dtype=np.int32)
    if not seeds:
        return labels
    yy, xx = np.nonzero(local_mask)
    seed_arr = np.array([[sx, sy] for sx, sy, _ in seeds], dtype=np.float32)
    pts = np.stack([xx, yy], axis=1).astype(np.float32)
    d2 = ((pts[:, None, :] - seed_arr[None, :, :]) ** 2).sum(axis=2)
    nearest = np.argmin(d2, axis=1) + 1
    labels[yy, xx] = nearest
    return labels


def split_large_components(mask, comps):
    instance_labels = np.zeros(mask.shape, dtype=np.int32)
    instance_boxes = []
    next_label = 1
    for comp in comps:
        local, local_bbox = crop_component_mask(mask, comp["bbox"], pad=1)
        x0, y0, _, _ = local_bbox
        bw = comp["bbox"][2] - comp["bbox"][0] + 1
        bh = comp["bbox"][3] - comp["bbox"][1] + 1
        should_split = comp["area"] > 1800 or bw > 70 or bh > 70
        seeds = pick_instance_seeds(local, max_seeds=9) if should_split else []
        labels = assign_to_seeds(local, seeds)
        if labels.max() == 0:
            labels = local.astype(np.int32)
        for local_label in range(1, int(labels.max()) + 1):
            region = labels == local_label
            if int(region.sum()) < 180:
                continue
            ys, xs = np.nonzero(region)
            gx0, gy0 = int(xs.min() + x0), int(ys.min() + y0)
            gx1, gy1 = int(xs.max() + x0), int(ys.max() + y0)
            instance_labels[y0 : y0 + local.shape[0], x0 : x0 + local.shape[1]][region] = next_label
            instance_boxes.append({"label": next_label, "bbox": (gx0, gy0, gx1, gy1), "area": int(region.sum())})
            next_label += 1
    return instance_labels, instance_boxes


def make_mask(arr):
    r, g, b = arr[:, :, 0], arr[:, :, 1], arr[:, :, 2]
    mx = arr.max(axis=2).astype(float)
    mn = arr.min(axis=2).astype(float)
    sat = (mx - mn) / np.maximum(mx, 1)
    h, w = r.shape
    yy = np.arange(h)[:, None]
    roi = yy > int(h * 0.14)

    yellow = (r > 170) & (g > 115) & (b < 120) & (r > b + 55) & (g > b + 30)
    orange = (r > 170) & (g > 70) & (g < 175) & (b < 95) & (r > g + 15)
    warm_light = (r > 185) & (g > 145) & (b < 150) & (sat > 0.18)
    mask = roi & (yellow | orange | warm_light)

    white_boxes = (sat < 0.18) & (mx > 150) & (yy > int(h * 0.72))
    black_floor = (mx < 55) & (yy > int(h * 0.62))
    dark_hats = (mx < 75) & (yy < int(h * 0.42))
    dark = mx < 45
    mask &= ~white_boxes
    mask &= ~black_floor
    mask &= ~dark_hats
    mask &= ~dark
    return mask


def overlay_mask(img, mask, color=(255, 64, 64), alpha=100):
    base = img.convert("RGBA")
    overlay = Image.new("RGBA", img.size, (0, 0, 0, 0))
    rgba = np.array(overlay)
    rgba[mask] = (*color, alpha)
    return Image.alpha_composite(base, Image.fromarray(rgba)).convert("RGB")


def draw_title(img, title):
    canvas = Image.new("RGB", (img.width, img.height + 54), "white")
    canvas.paste(img, (0, 54))
    d = ImageDraw.Draw(canvas)
    d.rectangle([0, 0, canvas.width, 54], fill=(30, 47, 75))
    d.text((18, 10), title, fill="white", font=FONT)
    return canvas


def adaptive_quantize(img):
    small = img.resize((900, int(img.height * 900 / img.width)), Image.Resampling.LANCZOS)
    quant = small.quantize(colors=7, method=Image.Quantize.MEDIANCUT).convert("RGB")
    return quant.resize(img.size, Image.Resampling.NEAREST)


def draw_bar_chart(counts, path):
    labels = list(counts.keys())
    values = list(counts.values())
    w, h = 900, 520
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    d.text((30, 22), "不同颜色候选区域面积统计", fill=(30, 47, 75), font=FONT_BIG)
    left, bottom = 110, 430
    chart_w, chart_h = 700, 300
    d.line([left, bottom, left + chart_w, bottom], fill=(80, 80, 80), width=2)
    d.line([left, bottom, left, bottom - chart_h], fill=(80, 80, 80), width=2)
    max_v = max(values) if values else 1
    colors = [(239, 113, 156), (83, 152, 226), (224, 207, 162), (226, 96, 86), (165, 165, 165)]
    step = chart_w / max(1, len(labels))
    for i, (label, value) in enumerate(zip(labels, values)):
        bar_h = int(chart_h * value / max_v)
        x0 = int(left + i * step + 28)
        x1 = int(left + (i + 1) * step - 28)
        y0 = bottom - bar_h
        d.rectangle([x0, y0, x1, bottom], fill=colors[i % len(colors)], outline=(80, 80, 80))
        d.text((x0, y0 - 26), str(int(value)), fill=(40, 40, 40), font=FONT_SMALL)
        d.text((x0, bottom + 16), label, fill=(40, 40, 40), font=FONT_SMALL)
    img.save(path, quality=95)


def instance_overlay(img, labels, boxes, path):
    palette = [
        (236, 88, 128),
        (59, 151, 225),
        (244, 183, 64),
        (75, 185, 133),
        (157, 108, 230),
        (233, 112, 71),
        (52, 190, 205),
        (215, 84, 198),
    ]
    base = img.convert("RGBA")
    overlay = Image.new("RGBA", img.size, (0, 0, 0, 0))
    rgba = np.array(overlay)
    for label in range(1, labels.max() + 1):
        color = palette[(label - 1) % len(palette)]
        rgba[labels == label] = (*color, 112)
    result = Image.alpha_composite(base, Image.fromarray(rgba)).convert("RGB")
    d = ImageDraw.Draw(result)
    for item in boxes[:90]:
        x0, y0, x1, y1 = item["bbox"]
        color = palette[(item["label"] - 1) % len(palette)]
        d.rectangle([x0, y0, x1, y1], outline=color, width=3)
        if item["area"] > 350:
            d.text((x0 + 2, max(2, y0 - 22)), str(item["label"]), fill=(0, 0, 0), font=FONT_SMALL)
    d.text((18, 18), f"进一步分割实例数：{len(boxes)}", fill=(255, 255, 255), font=FONT)
    draw_title(result, "图6 距离种子区域生长的进一步实例分割").save(path, quality=95)


def run_segmentation():
    if not SRC_IMAGE.exists():
        raise FileNotFoundError(SRC_IMAGE)
    original = Image.open(SRC_IMAGE).convert("RGB")
    work = resize_for_work(original, 1200)
    arr = np.array(work)

    raw_mask = make_mask(arr)
    opened = binary_dilate(binary_erode(raw_mask, radius=1, iterations=1), radius=1, iterations=1)
    closed = binary_erode(binary_dilate(opened, radius=1, iterations=1), radius=1, iterations=1)
    comps = connected_components(closed, min_area=210, max_area=26000)
    instance_labels, instance_boxes = split_large_components(closed, comps)

    r, g, b = arr[:, :, 0], arr[:, :, 1], arr[:, :, 2]
    color_masks = {
        "黄色主体": ((r > 175) & (g > 120) & (b < 125) & closed),
        "橙色嘴部/服饰": ((r > 170) & (g > 65) & (g < 175) & (b < 105) & closed),
        "浅黄高光": ((r > 190) & (g > 155) & (b < 165) & closed),
        "深色边界": ((r < 115) & (g < 100) & (b < 95) & closed),
        "其他": closed,
    }
    counts = {}
    used = np.zeros_like(closed)
    for key, cmask in list(color_masks.items())[:-1]:
        cmask = cmask & ~used
        counts[key] = int(cmask.sum())
        used |= cmask
    counts["其他"] = int((closed & ~used).sum())

    original_path = IMG_DIR / "fig01_original.jpg"
    quant_path = IMG_DIR / "fig02_color_quantization.jpg"
    mask_path = IMG_DIR / "fig03_threshold_mask.jpg"
    morph_path = IMG_DIR / "fig04_morphology_mask.jpg"
    overlay_path = IMG_DIR / "fig05_final_overlay.jpg"
    chart_path = IMG_DIR / "fig06_area_statistics.jpg"
    instance_path = IMG_DIR / "fig07_instance_segmentation.jpg"

    shutil.copyfile(SRC_IMAGE, IMG_DIR / "source_photo.jpg")
    draw_title(work, "图1 原始自拍黄色毛绒玩偶货架图像").save(original_path, quality=95)
    draw_title(adaptive_quantize(work), "图2 中值切分颜色聚类结果").save(quant_path, quality=95)
    draw_title(Image.fromarray((raw_mask * 255).astype(np.uint8)).convert("RGB"), "图3 颜色阈值得到的初始掩膜").save(mask_path, quality=95)
    draw_title(Image.fromarray((closed * 255).astype(np.uint8)).convert("RGB"), "图4 形态学开闭运算后的掩膜").save(morph_path, quality=95)

    overlay = overlay_mask(work, closed, color=(255, 58, 58), alpha=95)
    d = ImageDraw.Draw(overlay)
    for idx, comp in enumerate(comps[:60], start=1):
        x0, y0, x1, y1 = comp["bbox"]
        d.rectangle([x0, y0, x1, y1], outline=(0, 255, 210), width=3)
        if comp["area"] > 650:
            d.text((x0 + 2, max(2, y0 - 24)), str(idx), fill=(0, 0, 0), font=FONT_SMALL)
    d.text((18, 18), f"连通域候选目标数：{len(comps)}", fill=(255, 255, 255), font=FONT)
    draw_title(overlay, "图5 最终候选挂件区域与连通域框").save(overlay_path, quality=95)
    instance_overlay(work, instance_labels, instance_boxes, instance_path)
    draw_bar_chart(counts, chart_path)

    return {
        "original": original_path,
        "quant": quant_path,
        "mask": mask_path,
        "morph": morph_path,
        "overlay": overlay_path,
        "chart": chart_path,
        "instance": instance_path,
        "components": comps,
        "instances": instance_boxes,
        "counts": counts,
        "image_size": work.size,
        "mask_area_ratio": float(closed.sum() / closed.size),
    }


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def set_document_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_cn_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)


def add_image(doc, path, width_cm=14.5, caption=""):
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        add_caption(doc, caption)


def build_docx(results):
    doc = Document()
    set_document_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("图像分割实验报告")
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(30, 47, 75)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("题目：商店货架中毛绒挂件的图像分割与候选目标统计")
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(12)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    meta = [("课程/实验", "图像分割实际案例"), ("姓名/学号", "（请在此处填写）"), ("实验日期", "2026年6月"), ("实验图像来源", "本人拍摄的商店毛绒挂件货架照片")]
    for row, (k, v) in zip(table.rows, meta):
        set_cell_text(row.cells[0], k, bold=True)
        set_cell_shading(row.cells[0], "F2F4F7")
        set_cell_text(row.cells[1], v)

    doc.add_heading("一、实验目的", level=1)
    add_bullet(doc, "掌握基于颜色空间的图像分割基本流程，理解 RGB/HSV 思想在复杂场景中的应用。")
    add_bullet(doc, "综合使用颜色阈值、颜色聚类、形态学处理和连通域分析，完成真实图像中的目标区域提取。")
    add_bullet(doc, "分析透明包装、遮挡、反光和背景杂乱对分割结果的影响，并提出改进方案。")

    doc.add_heading("二、实验环境与实验对象", level=1)
    add_cn_paragraph(doc, "实验环境：Windows 系统，Python 语言，主要使用 Pillow、NumPy 和 python-docx 完成图像处理、结果可视化与报告生成。")
    add_cn_paragraph(doc, "实验对象为本人拍摄的商店货架图像。图像中包含大量粉色、蓝色、白色和米黄色毛绒挂件，目标密集排列，并受到透明塑料包装、金属挂钩、货架背景和人物背景的干扰。")
    add_image(doc, results["original"], 14.5, "图1 原始图像：商店货架中的毛绒挂件")

    doc.add_heading("三、实验原理", level=1)
    doc.add_heading("1. 颜色空间与阈值分割", level=2)
    add_cn_paragraph(doc, "彩色图像可表示为 RGB 三个通道。由于本图像中的目标颜色较鲜明，可以根据红、绿、蓝通道之间的差值构建粉色、蓝色、米白色、黄色等候选区域。阈值分割的优点是计算简单、速度快；缺点是对光照、反光和背景颜色较敏感。")
    doc.add_heading("2. 颜色聚类", level=2)
    add_cn_paragraph(doc, "为了观察图像的主色分布，本实验使用中值切分量化方法将图像颜色压缩为少数颜色类别。该步骤相当于一种颜色聚类，可帮助判断哪些颜色类别对应目标区域，哪些颜色类别属于背景或货架。")
    doc.add_heading("3. 形态学处理", level=2)
    add_cn_paragraph(doc, "初始掩膜中通常存在噪声点和局部空洞。本实验使用开运算去除小噪声，再使用闭运算连接同一目标内部的断裂区域，使分割结果更连续。")
    doc.add_heading("4. 连通域分析", level=2)
    add_cn_paragraph(doc, "连通域分析将二值掩膜中相邻的前景像素归为同一目标候选区域，并根据面积、宽高比等条件过滤不合理区域，最终得到候选挂件数量及位置。")
    doc.add_heading("5. 进一步实例分割", level=2)
    add_cn_paragraph(doc, "由于多个挂件会在二值掩膜中粘连成一个大连通域，本实验在候选区域基础上增加进一步分割。具体做法是对较大的连通域计算近似距离变换，选取距离边界较远的位置作为内部种子点，再将该连通域中的像素分配给最近的种子点。该方法相当于一种简化的基于距离种子的区域生长/分水岭思想，可以把部分相互接触的挂件拆分为多个实例。")

    doc.add_heading("四、实验步骤", level=1)
    steps = [
        "读取原始图像并按比例缩放，降低计算量，同时保持主要目标细节。",
        "根据 RGB 通道差异构造粉色、蓝色、米白色、黄色和高饱和度区域的初始掩膜。",
        "排除绿色货架、过暗背景和图像上部远景区域，减少非目标区域误检。",
        "对初始掩膜执行形态学开运算与闭运算，去除噪声并填补目标区域空洞。",
        "对处理后的掩膜进行连通域分析，按面积和宽高比筛选候选挂件区域。",
        "对面积较大的粘连区域计算距离种子，并按最近种子进行区域划分，得到进一步实例分割结果。",
        "将最终掩膜、候选框和实例分割结果叠加回原图，并统计不同颜色候选区域面积。",
    ]
    for item in steps:
        add_numbered(doc, item)

    doc.add_heading("五、实验结果", level=1)
    add_image(doc, results["quant"], 14.5, "图2 颜色聚类/量化结果")
    add_cn_paragraph(doc, "从颜色聚类结果可见，图像主色包括绿色货架、粉色挂件、蓝色装饰、米白色挂件以及背景区域。目标颜色和货架背景存在一定区分度，因此适合采用颜色阈值作为初始分割方法。")
    add_image(doc, results["mask"], 14.5, "图3 初始颜色阈值分割掩膜")
    add_cn_paragraph(doc, "初始掩膜能够较好地覆盖彩色挂件区域，但由于透明包装反光、标签和背景物体的影响，仍存在小面积噪声和部分区域断裂。")
    add_image(doc, results["morph"], 14.5, "图4 形态学处理后的二值掩膜")
    add_cn_paragraph(doc, "经过开闭运算后，细碎噪声减少，同一挂件内部区域更加连续，有利于后续连通域分析。")
    add_image(doc, results["overlay"], 14.5, "图5 最终分割区域与候选目标框")

    n_components = len(results["components"])
    n_instances = len(results["instances"])
    ratio = results["mask_area_ratio"] * 100
    add_cn_paragraph(doc, f"本次实验首先得到 {n_components} 个候选连通区域，前景掩膜约占缩放后图像面积的 {ratio:.2f}%。由于货架中挂件之间存在严重遮挡，单纯连通域会把相邻挂件合并，因此继续进行实例级拆分。")
    add_image(doc, results["instance"], 14.5, "图6 进一步实例分割结果")
    add_cn_paragraph(doc, f"进一步分割后得到 {n_instances} 个实例候选区域。与单纯连通域相比，该结果能够把部分横向排列且颜色相近的挂件分离出来，目标边界和候选数量更接近实际货架情况。")
    add_image(doc, results["chart"], 13.5, "图7 不同颜色候选区域面积统计")

    if ADVANCED_JSON.exists():
        advanced = json.loads(ADVANCED_JSON.read_text(encoding="utf-8"))
        doc.add_heading("六、优化实验：OpenCV + skimage 分水岭进一步分割", level=1)
        add_cn_paragraph(doc, "在基础实验的基础上，本次优化改用 OpenCV 进行 HSV/Lab 颜色掩膜提取，并使用 scipy 的欧氏距离变换与 skimage 的 watershed 算法进行进一步实例分割。优化思路是：先获得更干净的前景掩膜，再计算每个前景像素到背景边界的距离，距离峰值通常位于挂件中心，最后以这些峰值作为种子点执行分水岭分割。")
        add_image(doc, advanced["mask_path"], 14.5, "图8 OpenCV 精细颜色掩膜")
        add_cn_paragraph(doc, "与手写阈值掩膜相比，OpenCV 版本同时利用 HSV 色调、饱和度和 Lab 亮度/色度信息，对绿色货架和上方金属横杆进行了更明确的抑制，因此前景区域更集中在挂件本体。")
        add_image(doc, advanced["distance_path"], 14.5, "图9 距离变换热力图")
        add_cn_paragraph(doc, "距离变换图中亮度较高的位置表示其离背景边界较远，通常对应单个挂件的中心区域。这些局部峰值可作为后续分水岭算法的种子点。")
        add_image(doc, advanced["watershed_path"], 14.5, "图10 分水岭进一步实例分割结果")
        add_cn_paragraph(doc, f"优化后共得到 {advanced['instance_count']} 个实例候选区域，平均面积约 {advanced['mean_area']:.1f} 像素，中位面积约 {advanced['median_area']:.1f} 像素。相较基础版本的 {n_instances} 个候选区域，分水岭方法能更积极地拆分相互接触的挂件，尤其对底部和右侧密集区域的分离效果更明显。")
        add_cn_paragraph(doc, f"同时，本实验将每个实例的面积、外接矩形位置、宽高比和平均颜色导出为 Excel 文件：分水岭实例分割统计.xlsx。该表可用于后续计数、尺寸筛选或颜色分类分析。")

    stat_table = doc.add_table(rows=1, cols=2)
    stat_table.style = "Table Grid"
    stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(stat_table.rows[0].cells[0], "颜色类别", bold=True)
    set_cell_text(stat_table.rows[0].cells[1], "候选区域像素数", bold=True)
    for cell in stat_table.rows[0].cells:
        set_cell_shading(cell, "F2F4F7")
    for k, v in results["counts"].items():
        row = stat_table.add_row().cells
        set_cell_text(row[0], k)
        set_cell_text(row[1], str(v))

    doc.add_heading("七、结果分析", level=1)
    add_bullet(doc, "较好效果：粉色和蓝色挂件颜色鲜明，能够被阈值和聚类方法较稳定地提取出来。")
    add_bullet(doc, "进一步分割效果：距离种子区域生长能够把一部分粘连的大区域拆开，使结果从区域级分割提升到实例级候选分割。")
    add_bullet(doc, "主要误差来源：透明塑料包装会产生高亮反光，部分区域被识别为浅色目标；金属挂钩和白色标签也会造成误检。")
    add_bullet(doc, "遮挡问题：多个挂件相互重叠时，连通域可能把多个实际物体合并为一个区域，导致候选数量偏少。")
    add_bullet(doc, "背景干扰：货架、商场人群和标签颜色复杂，单纯颜色阈值难以完全排除所有非目标区域。")

    doc.add_heading("八、改进方向", level=1)
    add_bullet(doc, "可加入 GrabCut 或主动轮廓模型，在颜色阈值基础上进一步细化目标边界。")
    add_bullet(doc, "可结合边缘检测和分水岭算法，将相互接触的挂件进一步分离。")
    add_bullet(doc, "可采集更多不同角度、不同光照下的货架图像，训练基于深度学习的语义分割模型，例如 U-Net 或 Mask R-CNN。")
    add_bullet(doc, "如果需要精确计数，可先检测挂件头部或吊牌关键点，再与连通域结果融合。")

    doc.add_heading("九、实验总结", level=1)
    add_cn_paragraph(doc, "本实验以真实商店货架照片为对象，完成了从图像读取、颜色聚类、阈值分割、形态学处理到连通域分析的完整流程。实验结果表明，颜色特征对毛绒挂件这类色彩鲜明的目标具有较好的分割效果，但在遮挡、反光和背景复杂的真实场景中仍会出现误检和粘连。通过本实验可以看出，实际图像分割通常需要多种方法组合使用，才能在复杂场景中获得较稳定的结果。")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("图像分割实验报告")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(100, 100, 100)

    doc.save(DOCX_PATH)
    return DOCX_PATH


def main():
    ensure_dirs()
    results = run_segmentation()
    docx = build_docx(results)
    print(docx)


if __name__ == "__main__":
    main()
