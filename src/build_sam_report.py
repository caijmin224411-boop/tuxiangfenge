from pathlib import Path
import json

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SAM_DIR = ROOT / "report" / "sam_results"
OUT_DOCX = ROOT / "report" / "图像分割实验报告_SAM改进版.docx"
STATS = json.loads((SAM_DIR / "sam_stats.json").read_text(encoding="utf-8"))


def fmt(run, font="宋体", size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def p(doc, text, style=None):
    para = doc.add_paragraph(style=style)
    r = para.add_run(text)
    fmt(r)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    r = para.add_run(text)
    fmt(r)


def img(doc, path, cap, width=14.2):
    doc.add_picture(str(path), width=Cm(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(cap)
    fmt(r, size=9, color="666666")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell(cell, text, bold=False):
    cell.text = ""
    r = cell.paragraphs[0].add_run(text)
    fmt(r, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def setup(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    for name, size in [("Heading 1", 16), ("Heading 2", 13)]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string("2E74B5")


def main():
    doc = Document()
    setup(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("图像分割实验报告：SAM 改进版")
    fmt(r, font="黑体", size=22, bold=True, color="1E2F4B")

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("题目", "黄色毛绒玩偶的 SAM 框+点提示实例分割"),
        ("姓名/学号", "（请在此处填写）"),
        ("模型", STATS["model"]),
        ("图像来源", "本人拍摄的商店黄色毛绒玩偶货架照片"),
    ]
    for row, (k, v) in zip(table.rows, rows):
        shade(row.cells[0], "F2F4F7")
        cell(row.cells[0], k, True)
        cell(row.cells[1], v)

    doc.add_heading("一、问题与改进思路", level=1)
    p(doc, "该图像中玩偶、货架和灯光都集中在黄橙色范围内，单纯按颜色阈值分割会把货架和目标混在一起，因此颜色法效果很差。")
    p(doc, "本版改用 Segment Anything Model（SAM）。实验采用人工框提示和正样本点提示：框用于限定目标范围，点用于告诉模型框中真正要分割的玩偶位置；随后只保留包含正点的连通区域，减少货架横板和背景误分割。")

    doc.add_heading("二、实验步骤", level=1)
    for item in [
        "读取自拍图像并缩放到统一宽度。",
        "为主要可见玩偶设置框提示，并在每个目标中心设置正样本点。",
        "调用 facebook/sam-vit-base 生成多候选掩膜。",
        "选择 SAM 预测置信度最高的掩膜，并裁剪到提示框内。",
        "保留包含正样本点的连通块，去除同框内背景误分割。",
        "叠加显示实例掩膜，导出面积、外接框和置信度统计表。",
    ]:
        bullet(doc, item)

    doc.add_heading("三、实验结果", level=1)
    img(doc, SAM_DIR / "sam_01_original.jpg", "图1 原始图像")
    img(doc, SAM_DIR / "sam_02_prompt_boxes.jpg", "图2 SAM 框提示")
    img(doc, SAM_DIR / "sam_03_instances.jpg", "图3 SAM 框+点提示实例分割结果")
    p(doc, f"SAM 最终得到 {STATS['instance_count']} 个实例候选区域，平均预测置信度约 {STATS['mean_score']:.3f}。相比颜色阈值和椭圆形状先验，SAM 对大玩偶身体、脸部和前排玩偶边界的贴合明显更好。")

    doc.add_heading("四、结果分析", level=1)
    bullet(doc, "优点：不依赖黄橙色阈值，能在同色货架背景下利用模型学到的物体边界进行分割。")
    bullet(doc, "优点：框+点提示可以控制分割目标，适合复杂场景中的交互式实例分割。")
    bullet(doc, "不足：当多个玩偶严重遮挡或提示框包含货架横板时，仍会出现少量背景残留，需要更精细的负点或更紧的提示框继续修正。")
    bullet(doc, "改进方向：可增加负样本点、逐个目标微调提示框，或使用 SAM2/更大模型提升边界质量。")

    doc.add_heading("五、总结", level=1)
    p(doc, "本实验说明，真实复杂场景中颜色分割很容易失败；SAM 通过提示式分割能显著改善同色背景下的目标分割效果。对于本图像，SAM 框+点提示比颜色阈值和传统分水岭方法更适合作为最终方案。")

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
