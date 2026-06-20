from pathlib import Path
import json

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RESULT_DIR = ROOT / "report" / "shape_seed_results"
OUT_DOCX = ROOT / "report" / "图像分割实验报告_改进版_结构种子.docx"
STATS_JSON = RESULT_DIR / "shape_seed_stats.json"


def cn_run(run, font="宋体", size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    cn_run(r)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    cn_run(r)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    cn_run(r)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    cn_run(r, size=9, color="666666")


def image(doc, path, width=14.2, cap=None):
    doc.add_picture(str(path), width=Cm(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap:
        caption(doc, cap)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    cn_run(r, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def setup(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5")]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def build():
    stats = json.loads(STATS_JSON.read_text(encoding="utf-8"))
    doc = Document()
    setup(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("图像分割实验报告（改进版）")
    cn_run(r, font="黑体", size=22, bold=True, color="1E2F4B")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("题目：黄色毛绒玩偶的结构种子实例分割")
    cn_run(r, size=12)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    meta = [
        ("课程/实验", "图像分割实际案例"),
        ("姓名/学号", "（请在此处填写）"),
        ("实验日期", "2026年6月"),
        ("图像来源", "本人拍摄的商店黄色毛绒玩偶货架照片"),
    ]
    for row, (k, v) in zip(table.rows, meta):
        shade(row.cells[0], "F2F4F7")
        cell_text(row.cells[0], k, True)
        cell_text(row.cells[1], v)

    doc.add_heading("一、问题说明", level=1)
    para(doc, "原先使用颜色阈值区分黄色玩偶，效果较差。原因是本图像中玩偶、货架、灯光和部分背景都属于黄橙色系，颜色空间中的前景与背景高度重叠，单纯依靠颜色会把货架、灯光反光和玩偶主体混在一起。")
    para(doc, "因此，本改进版不再把颜色作为主要分割依据，而采用“结构特征种子 + 形状先验 + 局部区域竞争”的交互式实例分割方法。颜色只用于非常弱的目标范围约束和结果可视化。")

    doc.add_heading("二、实验方法", level=1)
    bullet(doc, "暗斑结构检测：在灰度图中提取眼睛、嘴巴等深色局部结构，这些结构比黄色颜色本身更能代表玩偶位置。")
    bullet(doc, "交互式种子修正：对主要可见玩偶手动补充中心种子，避免背景挂件和包装盒误入最终结果。")
    bullet(doc, "形状先验：每个种子生成一个局部椭圆候选域，模拟毛绒玩偶头身近似椭圆的外形。")
    bullet(doc, "局部竞争分割：候选域重叠时，将像素分配给归一化距离最近的种子，形成实例候选区域。")
    bullet(doc, "统计输出：计算每个实例的面积、外接矩形、宽高比和平均颜色，并导出 Excel 表格。")

    doc.add_heading("三、实验步骤", level=1)
    for step in [
        "读取原始图像并缩放到统一宽度，便于后续处理。",
        "在灰度图中检测深色结构斑点，获得眼睛、嘴巴等候选关键点。",
        "根据主要可见玩偶位置设置交互式种子，排除背景吊饰、包装盒和远处小物件。",
        "围绕每个种子建立椭圆形状先验，并在重叠区域进行最近种子竞争。",
        "输出实例分割叠加图，并导出实例统计表。",
    ]:
        numbered(doc, step)

    doc.add_heading("四、实验结果", level=1)
    image(doc, RESULT_DIR / "shape_01_original.jpg", 13.5, "图1 原始图像")
    image(doc, RESULT_DIR / "shape_03_seed_debug.jpg", 13.5, "图2 结构关键点与交互式种子")
    para(doc, "图2 中青色圈表示自动检测到的暗斑结构，红点表示最终用于分割的交互式种子。可以看到，种子主要落在前景主要玩偶上，避免了背景吊饰和包装盒参与最终实例分割。")
    image(doc, RESULT_DIR / "shape_04_shape_prior.jpg", 13.5, "图3 结构种子形状先验")
    image(doc, RESULT_DIR / "shape_06_instances.jpg", 13.5, "图4 改进后的实例分割结果")

    para(doc, f"最终保留 {stats['instance_count']} 个主要实例候选区域，平均面积约 {stats['mean_area']:.1f} 像素，中位面积约 {stats['median_area']:.1f} 像素。与颜色阈值法相比，该结果更符合前景玩偶的实际位置，背景货架大面积误分割明显减少。")

    doc.add_heading("五、结果分析", level=1)
    bullet(doc, "优点：不依赖黄色/橙色阈值，能避开同色货架和暖色灯光造成的大面积误检。")
    bullet(doc, "优点：交互式种子使目标范围更明确，适合复杂真实场景中的实例分割实验。")
    bullet(doc, "不足：椭圆形状先验不能精确贴合玩偶边界，边缘仍是近似结果。")
    bullet(doc, "不足：部分相互遮挡的玩偶仍可能被合并，背景中的远处玩偶没有全部纳入统计。")

    doc.add_heading("六、总结", level=1)
    para(doc, "本实验表明，在同色背景严重的真实图像中，单纯颜色分割并不可靠。改进后的结构种子方法利用眼睛、嘴巴等局部结构确定实例中心，再结合形状先验完成区域竞争分割，结果更稳定，也更符合图像分割实际案例的要求。若继续优化，可使用 SAM、Mask R-CNN 或人工精细标注数据进行深度学习实例分割。")

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
